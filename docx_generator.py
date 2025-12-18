"""
DOCX Resume Generator for ATS Compatibility

Generates Word documents (.docx) optimized for Applicant Tracking Systems.
Uses python-docx library for clean, ATS-friendly formatting.
"""

from pathlib import Path
from typing import Dict, Any
import re

# Import config and contact info
from config import get_contact_info
from resume_parser import parse_markdown_resume, ResumeData
CONTACT_INFO = get_contact_info()

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Install with: pip install python-docx")


def parse_markdown_to_docx_data(markdown_text: str) -> Dict[str, Any]:
    """
    Parse markdown resume into structured data for DOCX generation.

    Uses the shared resume_parser module and converts to dict format.
    """
    resume_data = parse_markdown_resume(markdown_text)

    # Convert ResumeData to dict format expected by generate_docx_resume
    return {
        'name': resume_data.name.upper(),
        'title': resume_data.title,
        'contact_info': resume_data.contact_info,
        'summary': resume_data.summary,
        'experience': [
            {
                'company': job.company,
                'title': job.title,
                'dates': job.dates,
                'bullets': job.bullets
            }
            for job in resume_data.experience
        ],
        'achievements': resume_data.achievements,
        'skills': [f"{cat}: {skills}" for cat, skills in resume_data.skills.items()],
        'education': resume_data.education,
        'certifications': resume_data.certifications
    }


def generate_docx_resume(resume_data: Dict[str, Any], output_path: Path) -> Path:
    """
    Generate ATS-friendly Word document resume.

    Args:
        resume_data: Dictionary with resume sections (name, experience, skills, etc.)
        output_path: Path to save .docx file

    Returns:
        Path to generated .docx file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document()

    # Set document margins (0.5 inch all around for ATS compatibility)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Header: Name (use hardcoded contact info)
    name = doc.add_paragraph()
    name_run = name.add_run(CONTACT_INFO['name'].upper())
    name_run.font.size = Pt(18)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header: Title
    title = doc.add_paragraph()
    title_run = title.add_run(resume_data.get('title', 'Senior Product Manager'))
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue accent
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Header: Contact Info (use hardcoded contact info)
    contact_text = f"{CONTACT_INFO['phone']} | {CONTACT_INFO['email']} | {CONTACT_INFO['linkedin']} | {CONTACT_INFO['location']}"
    contact = doc.add_paragraph(contact_text)
    contact.runs[0].font.size = Pt(10)
    contact.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add spacing after header
    doc.add_paragraph()

    # Professional Summary
    if resume_data.get('summary'):
        add_section_header(doc, 'PROFESSIONAL SUMMARY')
        summary_para = doc.add_paragraph(resume_data['summary'])
        summary_para.runs[0].font.size = Pt(10)
        doc.add_paragraph()

    # Experience Section
    if resume_data.get('experience'):
        add_section_header(doc, 'EXPERIENCE')

        for job in resume_data['experience']:
            # Company | Title | Dates
            job_header = doc.add_paragraph()
            company_run = job_header.add_run(f"{job.get('company', '')} | ")
            company_run.font.bold = True
            company_run.font.size = Pt(11)

            title_run = job_header.add_run(f"{job.get('title', '')} | ")
            title_run.font.size = Pt(11)

            dates_run = job_header.add_run(job.get('dates', ''))
            dates_run.font.size = Pt(10)
            dates_run.font.color.rgb = RGBColor(102, 102, 102)  # Gray

            # Bullets
            for bullet in job.get('bullets', []):
                bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                bullet_para.runs[0].font.size = Pt(10)

            # Add spacing between jobs
            doc.add_paragraph()

    # Key Achievements Section
    if resume_data.get('achievements'):
        add_section_header(doc, 'KEY ACHIEVEMENTS')

        for achievement in resume_data['achievements']:
            if isinstance(achievement, dict):
                # Achievement with title and description
                ach_para = doc.add_paragraph()
                title_run = ach_para.add_run(f"{achievement.get('title', '')}: ")
                title_run.font.bold = True
                title_run.font.size = Pt(10)

                desc_run = ach_para.add_run(achievement.get('description', ''))
                desc_run.font.size = Pt(10)
            else:
                # Simple achievement string
                ach_para = doc.add_paragraph(achievement, style='List Bullet')
                ach_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()

    # Skills Section
    if resume_data.get('skills'):
        add_section_header(doc, 'SKILLS')

        skills_para = doc.add_paragraph()
        skills_text = ', '.join(resume_data['skills']) if isinstance(resume_data['skills'], list) else resume_data['skills']
        skills_run = skills_para.add_run(skills_text)
        skills_run.font.size = Pt(10)

        doc.add_paragraph()

    # Education Section
    if resume_data.get('education'):
        add_section_header(doc, 'EDUCATION')

        edu = resume_data['education']
        if isinstance(edu, dict):
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(f"{edu.get('degree', '')}, ")
            degree_run.font.bold = True
            degree_run.font.size = Pt(10)

            school_run = edu_para.add_run(f"{edu.get('school', '')}, ")
            school_run.font.size = Pt(10)

            dates_run = edu_para.add_run(edu.get('dates', ''))
            dates_run.font.size = Pt(10)
        else:
            edu_para = doc.add_paragraph(str(edu))
            edu_para.runs[0].font.size = Pt(10)

        doc.add_paragraph()

    # Certifications Section
    if resume_data.get('certifications'):
        add_section_header(doc, 'CERTIFICATIONS')

        for cert in resume_data['certifications']:
            if isinstance(cert, dict):
                cert_para = doc.add_paragraph()
                cert_run = cert_para.add_run(cert.get('title', ''))
                cert_run.font.bold = True
                cert_run.font.size = Pt(10)

                if cert.get('description'):
                    desc_para = doc.add_paragraph(cert['description'])
                    desc_para.runs[0].font.size = Pt(9)
            else:
                cert_para = doc.add_paragraph(cert, style='List Bullet')
                cert_para.runs[0].font.size = Pt(10)

    # Save document
    doc.save(str(output_path))
    return output_path


def add_section_header(doc: Document, text: str):
    """Add a formatted section header to the document."""
    header = doc.add_paragraph()
    header_run = header.add_run(text)
    header_run.font.size = Pt(12)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0, 0, 0)


def generate_docx_cover_letter(markdown_text: str, output_path: Path) -> Path:
    """
    Generate ATS-friendly Word document cover letter from markdown.

    Args:
        markdown_text: Markdown formatted cover letter text
        output_path: Path to save .docx file

    Returns:
        Path to generated .docx file
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Parse the markdown text
    lines = markdown_text.split('\n')

    for i, line in enumerate(lines):
        line_stripped = line.strip()

        # Skip empty lines (but add spacing)
        if not line_stripped:
            if i > 0:  # Don't add space at the beginning
                doc.add_paragraph()
            continue

        # Headers (##)
        if line_stripped.startswith('## '):
            header_text = line_stripped.replace('## ', '')
            para = doc.add_paragraph()
            run = para.add_run(header_text)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235)  # Blue accent
            continue

        # Bold text (**text**)
        if '**' in line_stripped:
            para = doc.add_paragraph()
            # Simple bold handling
            parts = line_stripped.split('**')
            for j, part in enumerate(parts):
                run = para.add_run(part)
                if j % 2 == 1:  # Odd indices are bold
                    run.font.bold = True
                run.font.size = Pt(11)
            continue

        # Bullet points (lines starting with -)
        if line_stripped.startswith('- '):
            text = line_stripped[2:]  # Remove '- '
            para = doc.add_paragraph(text, style='List Bullet')
            para.runs[0].font.size = Pt(11)
            continue

        # Regular paragraphs
        para = doc.add_paragraph(line_stripped)
        para.runs[0].font.size = Pt(11)

    # Save the document
    doc.save(str(output_path))
    return output_path
